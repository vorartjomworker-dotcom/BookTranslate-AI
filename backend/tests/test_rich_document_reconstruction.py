import base64
from io import BytesIO
from pathlib import Path

from docx import Document

from app.services.docx_parser import parse_docx
from app.services.reconstruction import reconstruct_docx


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nGQAAAAASUVORK5CYII="
)


def _counts(document):
    result = {}
    for chapter in document.chapters:
        for block in chapter.blocks:
            result[block.block_type] = result.get(block.block_type, 0) + 1
    return result


def test_docx_rich_elements_survive_parse_reconstruct_parse(tmp_path: Path) -> None:
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "reconstructed.docx"

    source = Document()
    source.core_properties.title = "Round Trip"
    source.add_heading("Chapter One", level=1)
    source.add_heading("Section A", level=2)
    source.add_paragraph("Intro paragraph.")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    source.add_picture(BytesIO(_PNG_1X1))
    source.add_paragraph("Figure 1. Pixel", style="Caption")
    source.add_paragraph("First item", style="List Bullet")
    source.save(source_path)

    first = parse_docx(source_path)
    assert len(first.assets) == 1
    assert _counts(first)["table"] == 1
    assert _counts(first)["figure"] == 1
    assert _counts(first)["caption"] == 1

    reconstruct_docx(first, output_path)
    second = parse_docx(output_path)

    assert output_path.exists()
    assert len(second.assets) == 1
    assert _counts(second)["table"] == 1
    assert _counts(second)["figure"] == 1
    assert _counts(second)["caption"] == 1
    assert _counts(second)["list_item"] == 1
    assert second.chapters[0].title == "Chapter One"
    assert second.chapters[0].sections[0].title == "Section A"
