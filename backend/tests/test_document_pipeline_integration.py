import asyncio
import base64
import os
import shutil
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from starlette.datastructures import UploadFile

from app.api.upload import upload_book
from app.core.config import settings
from app.db import AsyncSessionLocal, engine
from app.models.book import Book
from app.services.document_export import load_normalized_document
from app.services.docx_parser import parse_docx
from app.services.reconstruction import reconstruct_docx


pytestmark = pytest.mark.skipif(
    os.getenv("RUN_DB_INTEGRATION") != "1",
    reason="requires migrated PostgreSQL integration database",
)

_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nGQAAAAASUVORK5CYII="
)


def _block_counts(document) -> dict[str, int]:
    counts: dict[str, int] = {}
    for chapter in document.chapters:
        for block in chapter.blocks:
            counts[block.block_type] = counts.get(block.block_type, 0) + 1
    return counts


async def _round_trip() -> None:
    stream = BytesIO()
    source = Document()
    source.core_properties.title = "Database Round Trip"
    source.add_heading("Chapter One", level=1)
    source.add_heading("Section A", level=2)
    source.add_paragraph("Persistent paragraph.")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Latency"
    table.cell(1, 1).text = "10 us"
    source.add_picture(BytesIO(_PNG_1X1))
    source.add_paragraph("Figure 1. Test image", style="Caption")
    source.save(stream)
    stream.seek(0)

    upload = UploadFile(filename="database-round-trip.docx", file=stream)
    upload_root = Path(settings.upload_dir)
    output_path = upload_root / "integration" / "reconstructed.docx"
    stored_source: Path | None = None

    try:
        async with AsyncSessionLocal() as db:
            response = await upload_book(
                file=upload,
                title=None,
                source_language="en",
                target_language="ru",
                db=db,
            )
            assert response.chapters == 1
            assert response.sections == 1
            assert response.tables == 1
            assert response.figures == 1
            assert response.captions == 1
            assert response.assets == 1

            normalized = await load_normalized_document(db, response.book_id, upload_root)
            assert normalized is not None
            persisted_counts = _block_counts(normalized)
            assert persisted_counts["table"] == 1
            assert persisted_counts["figure"] == 1
            assert persisted_counts["caption"] == 1

            reconstruct_docx(normalized, output_path)
            reparsed = parse_docx(output_path)
            reparsed_counts = _block_counts(reparsed)
            assert reparsed_counts["table"] == 1
            assert reparsed_counts["figure"] == 1
            assert reparsed_counts["caption"] == 1
            assert reparsed.chapters[0].title == "Chapter One"
            assert reparsed.chapters[0].sections[0].title == "Section A"

            book = await db.get(Book, response.book_id)
            stored_source = upload_root / book.stored_filename if book and book.stored_filename else None
            await db.delete(book)
            await db.commit()
    finally:
        if stored_source is not None:
            stored_source.unlink(missing_ok=True)
        if "response" in locals():
            shutil.rmtree(upload_root / "assets" / str(response.book_id), ignore_errors=True)
        output_path.unlink(missing_ok=True)
        await engine.dispose()


def test_upload_database_reconstruction_round_trip() -> None:
    asyncio.run(_round_trip())
