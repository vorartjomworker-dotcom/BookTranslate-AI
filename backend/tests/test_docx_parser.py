from pathlib import Path

from docx import Document

from app.services.docx_parser import parse_docx


def test_docx_parser_splits_heading_one_into_chapters(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.core_properties.title = "Sample Book"
    document.add_heading("Chapter One", level=1)
    document.add_paragraph("First paragraph.")
    document.add_heading("Chapter Two", level=1)
    document.add_paragraph("Second paragraph.")
    document.save(path)

    parsed = parse_docx(path)

    assert parsed.title == "Sample Book"
    assert [chapter.title for chapter in parsed.chapters] == ["Chapter One", "Chapter Two"]
    assert parsed.chapters[0].paragraphs == ["First paragraph."]
    assert parsed.chapters[1].paragraphs == ["Second paragraph."]


def test_docx_parser_preserves_sections_and_block_order(tmp_path: Path) -> None:
    path = tmp_path / "structured.docx"
    document = Document()
    document.add_heading("Chapter One", level=1)
    document.add_heading("Section A", level=2)
    document.add_paragraph("Paragraph A.")
    document.add_heading("Subsection A.1", level=3)
    document.add_paragraph("Quoted text.", style="Quote")
    document.save(path)

    chapter = parse_docx(path).chapters[0]

    assert [(section.level, section.title, section.parent_position) for section in chapter.sections] == [
        (2, "Section A", None),
        (3, "Subsection A.1", 0),
    ]
    assert [block.block_type for block in chapter.blocks] == [
        "heading",
        "paragraph",
        "heading",
        "blockquote",
    ]
    assert [block.section_position for block in chapter.blocks] == [0, 0, 1, 1]
