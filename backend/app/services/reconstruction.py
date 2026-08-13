from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

from app.services.document_parser import NormalizedDocument


def _add_paragraph(document: DocxDocument, text: str, style: str | None = None):
    if style:
        try:
            return document.add_paragraph(text, style=style)
        except KeyError:
            pass
    return document.add_paragraph(text)


def _append_hyperlink(paragraph, text: str, href: str) -> None:
    if not href:
        return
    if href.startswith("#"):
        paragraph.add_run(f" [{text or href}]")
        return
    relation_id = paragraph.part.relate_to(href, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text or href
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _apply_inline_fidelity(paragraph, metadata: dict) -> None:
    links = metadata.get("hyperlinks") or []
    if links:
        paragraph.add_run(" ")
        for index, link in enumerate(links):
            if index:
                paragraph.add_run(" · ")
            _append_hyperlink(paragraph, str(link.get("text") or link.get("href") or "link"), str(link.get("href") or ""))
    for reference in metadata.get("footnote_references") or []:
        paragraph.add_run(f" [fn:{reference}]")
    for xml in metadata.get("omml") or []:
        try:
            paragraph._p.append(parse_xml(xml))
        except Exception:
            paragraph.add_run(" [formula]")


def reconstruct_docx(source: NormalizedDocument, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    if source.title:
        document.core_properties.title = source.title

    assets = {asset.position: asset for asset in source.assets}

    for chapter in source.chapters:
        if chapter.title:
            document.add_heading(chapter.title, level=1)

        for block in sorted(chapter.blocks, key=lambda item: item.position):
            text = block.source_text or ""
            metadata = dict(block.metadata_json or {})
            paragraph = None
            if block.block_type == "heading":
                level = int(metadata.get("level", 2))
                paragraph = document.add_heading(text, level=max(1, min(level, 9)))
            elif block.block_type == "paragraph":
                paragraph = _add_paragraph(document, text)
            elif block.block_type == "list_item":
                style = "List Number" if metadata.get("list_kind") == "number" else "List Bullet"
                paragraph = _add_paragraph(document, text, style)
            elif block.block_type == "code":
                paragraph = _add_paragraph(document, text, "No Spacing")
            elif block.block_type == "blockquote":
                paragraph = _add_paragraph(document, text, "Quote")
            elif block.block_type == "caption":
                paragraph = _add_paragraph(document, text, "Caption")
            elif block.block_type == "table":
                cells = metadata.get("cells") or []
                rows_count = len(cells)
                columns_count = max((len(row) for row in cells), default=0)
                if rows_count and columns_count:
                    table = document.add_table(rows=rows_count, cols=columns_count)
                    for row_index, row in enumerate(cells):
                        for column_index, value in enumerate(row):
                            table.cell(row_index, column_index).text = str(value)
            elif block.block_type == "figure":
                asset_position = metadata.get("asset_position")
                asset = assets.get(asset_position) if isinstance(asset_position, int) else None
                if asset and asset.data:
                    document.add_picture(BytesIO(asset.data))
            if paragraph is not None:
                _apply_inline_fidelity(paragraph, metadata)

    document.save(output_path)
    return output_path
