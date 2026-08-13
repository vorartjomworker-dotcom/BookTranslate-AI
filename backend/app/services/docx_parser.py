import hashlib
import re
import zipfile
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from lxml import etree

from app.services.document_parser import (
    NormalizedAsset,
    NormalizedBlock,
    NormalizedChapter,
    NormalizedDocument,
    NormalizedSection,
)


_HEADING_LEVEL = re.compile(r"(?:heading|заголовок)\s*(\d+)", re.IGNORECASE)
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _heading_level(style_name: str) -> int | None:
    match = _HEADING_LEVEL.search(style_name)
    return int(match.group(1)) if match else None


def _block_type(style_name: str) -> str:
    lowered = style_name.lower()
    if "list" in lowered or "спис" in lowered:
        return "list_item"
    if "code" in lowered or "preformatted" in lowered or "код" in lowered:
        return "code"
    if "quote" in lowered or "цитат" in lowered:
        return "blockquote"
    return "paragraph"


def _is_caption_style(style_name: str) -> bool:
    lowered = style_name.lower()
    return "caption" in lowered or "подпис" in lowered or "название" in lowered


def _iter_body_items(document: DocxDocument) -> Iterator[Paragraph | DocxTable]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, document)


def _read_note_bodies(path: Path, note_type: str) -> dict[str, str]:
    part_name = f"word/{note_type}s.xml"
    element_name = note_type
    try:
        with zipfile.ZipFile(path) as archive:
            if part_name not in archive.namelist():
                return {}
            root = etree.fromstring(archive.read(part_name))
    except (KeyError, zipfile.BadZipFile, etree.XMLSyntaxError):
        return {}

    result: dict[str, str] = {}
    for note in root.findall(f"{{{_W_NS}}}{element_name}"):
        note_id = note.get(f"{{{_W_NS}}}id")
        if note_id is None:
            continue
        try:
            if int(note_id) <= 0:
                continue
        except ValueError:
            pass
        paragraphs: list[str] = []
        for paragraph in note.findall(f".//{{{_W_NS}}}p"):
            text = "".join(node.text or "" for node in paragraph.findall(f".//{{{_W_NS}}}t")).strip()
            if text:
                paragraphs.append(text)
        body = "\n".join(paragraphs).strip()
        if body:
            result[str(note_id)] = body
    return result


def _paragraph_inline_metadata(paragraph: Paragraph, document: DocxDocument) -> dict:
    hyperlinks = []
    for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
        relation_id = hyperlink.get(qn("r:id"))
        anchor = hyperlink.get(qn("w:anchor"))
        href = None
        if relation_id:
            relationship = document.part.rels.get(relation_id)
            if relationship is not None:
                href = getattr(relationship, "target_ref", None)
        if not href and anchor:
            href = f"#{anchor}"
        text = "".join(node.text or "" for node in hyperlink.xpath(".//w:t")).strip()
        if href:
            hyperlinks.append({"href": str(href), "text": text})

    footnote_references = []
    for reference in paragraph._p.xpath(".//w:footnoteReference"):
        value = reference.get(qn("w:id"))
        if value is not None:
            footnote_references.append(str(value))

    endnote_references = []
    for reference in paragraph._p.xpath(".//w:endnoteReference"):
        value = reference.get(qn("w:id"))
        if value is not None:
            endnote_references.append(str(value))

    omml = []
    for math in paragraph._p.xpath(".//m:oMath | .//m:oMathPara"):
        omml.append(etree.tostring(math, encoding="unicode"))

    metadata: dict = {}
    if hyperlinks:
        metadata["hyperlinks"] = hyperlinks
    if footnote_references:
        metadata["footnote_references"] = footnote_references
    if endnote_references:
        metadata["endnote_references"] = endnote_references
    if omml:
        metadata["omml"] = omml
    return metadata


def _register_asset(
    assets: list[NormalizedAsset],
    by_hash: dict[str, int],
    data: bytes,
    filename: str | None,
    media_type: str | None,
    metadata_json: dict,
) -> int:
    digest = hashlib.sha256(data).hexdigest()
    existing = by_hash.get(digest)
    if existing is not None:
        return existing

    position = len(assets)
    assets.append(
        NormalizedAsset(
            position=position,
            original_filename=filename,
            media_type=media_type,
            data=data,
            sha256=digest,
            metadata_json=metadata_json,
        )
    )
    by_hash[digest] = position
    return position


def _paragraph_images(
    paragraph: Paragraph,
    document: DocxDocument,
    assets: list[NormalizedAsset],
    by_hash: dict[str, int],
) -> list[tuple[int, str | None]]:
    results: list[tuple[int, str | None]] = []
    alt_text = None
    doc_props = paragraph._p.xpath(".//wp:docPr")
    if doc_props:
        alt_text = doc_props[0].get("descr") or doc_props[0].get("title")

    for blip in paragraph._p.xpath(".//a:blip"):
        relation_id = blip.get(qn("r:embed"))
        if not relation_id:
            continue
        image_part = document.part.related_parts.get(relation_id)
        if image_part is None or not hasattr(image_part, "blob"):
            continue
        filename = Path(str(getattr(image_part, "partname", "image"))).name
        media_type = getattr(image_part, "content_type", None)
        position = _register_asset(
            assets,
            by_hash,
            image_part.blob,
            filename,
            media_type,
            {"source": "docx", "relation_id": relation_id},
        )
        results.append((position, alt_text))
    return results


def _append_text_block(
    chapter: NormalizedChapter,
    text: str,
    style_name: str,
    section_position: int | None,
    last_target_block_position: int | None,
    inline_metadata: dict,
) -> int | None:
    if _is_caption_style(style_name):
        chapter.paragraphs.append(text)
        chapter.blocks.append(
            NormalizedBlock(
                position=len(chapter.blocks),
                block_type="caption",
                source_text=text,
                section_position=section_position,
                metadata_json={
                    "style": style_name,
                    "target_block_position": last_target_block_position,
                    **inline_metadata,
                },
            )
        )
        return last_target_block_position

    block_type = _block_type(style_name)
    metadata = {"style": style_name, **inline_metadata}
    if block_type == "list_item":
        lowered = style_name.lower()
        metadata["list_kind"] = "number" if "number" in lowered or "нумер" in lowered else "bullet"

    chapter.paragraphs.append(text)
    chapter.blocks.append(
        NormalizedBlock(
            position=len(chapter.blocks),
            block_type=block_type,
            source_text=text,
            section_position=section_position,
            metadata_json=metadata,
        )
    )
    return last_target_block_position


def _append_referenced_notes(
    chapters: list[NormalizedChapter],
    footnotes: dict[str, str],
    endnotes: dict[str, str],
) -> None:
    emitted: set[tuple[str, str]] = set()
    for chapter in chapters:
        references: list[tuple[str, str]] = []
        for block in chapter.blocks:
            metadata = block.metadata_json or {}
            references.extend(("footnote", str(value)) for value in metadata.get("footnote_references") or [])
            references.extend(("endnote", str(value)) for value in metadata.get("endnote_references") or [])
        for note_type, note_id in references:
            key = (note_type, note_id)
            if key in emitted:
                continue
            body = (footnotes if note_type == "footnote" else endnotes).get(note_id)
            if not body:
                continue
            chapter.blocks.append(
                NormalizedBlock(
                    position=len(chapter.blocks),
                    block_type=note_type,
                    source_text=body,
                    section_position=None,
                    metadata_json={"note_id": note_id, "note_type": note_type, "source": "docx"},
                )
            )
            emitted.add(key)


def parse_docx(path: Path) -> NormalizedDocument:
    document = Document(path)
    title = (document.core_properties.title or "").strip() or path.stem
    footnotes = _read_note_bodies(path, "footnote")
    endnotes = _read_note_bodies(path, "endnote")

    chapters: list[NormalizedChapter] = []
    assets: list[NormalizedAsset] = []
    assets_by_hash: dict[str, int] = {}
    current = NormalizedChapter(title=None)
    section_stack: dict[int, int] = {}
    current_section_position: int | None = None
    last_target_block_position: int | None = None

    for item in _iter_body_items(document):
        if isinstance(item, DocxTable):
            cells = [[cell.text.strip() for cell in row.cells] for row in item.rows]
            columns_count = max((len(row) for row in cells), default=0)
            source_text = "\n".join("\t".join(row) for row in cells)
            block_position = len(current.blocks)
            current.blocks.append(
                NormalizedBlock(
                    position=block_position,
                    block_type="table",
                    source_text=source_text,
                    section_position=current_section_position,
                    metadata_json={
                        "cells": cells,
                        "rows_count": len(cells),
                        "columns_count": columns_count,
                        "style": item.style.name if item.style else None,
                    },
                )
            )
            last_target_block_position = block_position
            continue

        paragraph = item
        text = paragraph.text.strip()
        style_name = (paragraph.style.name or "") if paragraph.style else ""
        heading_level = _heading_level(style_name)
        inline_metadata = _paragraph_inline_metadata(paragraph, document)

        if heading_level == 1 and text:
            if current.paragraphs or current.title or current.blocks:
                chapters.append(current)
            current = NormalizedChapter(title=text)
            section_stack = {}
            current_section_position = None
            last_target_block_position = None
            continue

        if heading_level and heading_level >= 2 and text:
            parent_position = None
            for candidate_level in range(heading_level - 1, 1, -1):
                if candidate_level in section_stack:
                    parent_position = section_stack[candidate_level]
                    break

            section_position = len(current.sections)
            current.sections.append(
                NormalizedSection(
                    position=section_position,
                    level=heading_level,
                    title=text,
                    parent_position=parent_position,
                    metadata_json={"style": style_name, **inline_metadata},
                )
            )
            current.blocks.append(
                NormalizedBlock(
                    position=len(current.blocks),
                    block_type="heading",
                    source_text=text,
                    section_position=section_position,
                    metadata_json={"level": heading_level, "style": style_name, **inline_metadata},
                )
            )
            section_stack = {level: pos for level, pos in section_stack.items() if level < heading_level}
            section_stack[heading_level] = section_position
            current_section_position = section_position
            continue

        if text:
            last_target_block_position = _append_text_block(
                current,
                text,
                style_name,
                current_section_position,
                last_target_block_position,
                inline_metadata,
            )

        for asset_position, alt_text in _paragraph_images(paragraph, document, assets, assets_by_hash):
            block_position = len(current.blocks)
            current.blocks.append(
                NormalizedBlock(
                    position=block_position,
                    block_type="figure",
                    source_text=None,
                    section_position=current_section_position,
                    metadata_json={
                        "asset_position": asset_position,
                        "alt_text": alt_text,
                        "source": "docx",
                    },
                )
            )
            last_target_block_position = block_position

    if current.paragraphs or current.title or current.blocks:
        chapters.append(current)

    if not chapters:
        chapters = [NormalizedChapter(title=title)]
    elif len(chapters) == 1 and chapters[0].title is None:
        chapters[0].title = title

    _append_referenced_notes(chapters, footnotes, endnotes)
    return NormalizedDocument(title=title, chapters=chapters, assets=assets)
